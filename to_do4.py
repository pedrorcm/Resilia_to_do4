import PyPDF2
import docx
import re

def recebendo_arquivo():
  n_arq = input('Digite o nome do arquivo (com extensão), na pasta, que deseja analisar: ')

  if '.pdf' in n_arq:
    ext = 'pdf'
    return n_arq, ext

  elif '.docx' in n_arq:
    ext = 'docx'
    return n_arq, ext

  else:
    print('Entre com um arquivo válido. ')
    exit()



def pdf(nome_arq):

  pdfF = open(nome_arq, 'rb')
  readr = PyPDF2.PdfReader(pdfF)

  print(f'O PDF carregado tem {readr.numPages} paginas')

  return readr, readr.numPages


def word_doc(nome_arq):
  docxF = docx.Document(nome_arq)
  parag = len(docxF.paragraphs)

  print(f'O docx carregado tem {parag} linhas')

  return docxF, parag


def abrindo():
  nome_arquivo, extensao = recebendo_arquivo()

  if extensao == 'pdf':
    read, nP = pdf(nome_arquivo)
    return read, nP, extensao


  elif extensao == 'docx':
    docxFile, paragr = word_doc(nome_arquivo)
    return docxFile, paragr, extensao



def get_all_text():
  rea, nPaginas, exte = abrindo()

  full_text = []


  if exte == 'pdf':
    for n in range(0,nPaginas):
      pagina = rea.getPage(n)
      full_text.append(str(pagina.extract_text()))
        
    full_text = str(full_text[0])

    return full_text

  if exte == 'docx':
    for para in rea.paragraphs:
      full_text.append(para.text)

    return str(full_text)


def treating_text(txt=get_all_text()):
  txt = re.sub("[^\w @/.]", "",txt).lower()

  return txt


def getting_email(txt):
  em = re.findall('[\w]{0,40}@[\w]{0,12}.[\w]{0,4}', txt)

  return em


def compatibilidade(cmpt, prof_elegiveis, mail):
  if len(cmpt) == 0 and len(prof_elegiveis) == 0:
    prof_elegiveis.add('Nenhuma no momento')
    cmpt.append('para as vagas abertas no momento não são compatíveis')


  candidato = (f'''
O candidato identificado pelo email: {mail}
está apto a se candidatar para as vagas: {" e ".join([i for i in prof_elegiveis])}.

Suas competências desejadas são {", ".join([a.capitalize() for a in cmpt])}
''')

  return candidato



def candidato():

  texto_str = treating_text()
  email = getting_email(texto_str)

  competencias, profissoes_elegiveis = acessando_chaves(texto_str)

  candidat = compatibilidade(competencias, profissoes_elegiveis, email)

  return candidat



def acessando_chaves(txt):
  dc = {'analista de dados':('python', 'powerbi', 'sql','comunicação'), 'cientista de dados': ('python', 'banco de dados', 'machine learning', 'resolução de problemas','estatística')}

  competencias = []
  prof_elegiveis = set()


  for parcv in dc.items():

    for elemento_do_valor in parcv[1]:
      if elemento_do_valor in txt:
        competencias.append(elemento_do_valor)
        prof_elegiveis.add(parcv[0])

  
  return competencias, prof_elegiveis



def main():
  print(candidato())


main()